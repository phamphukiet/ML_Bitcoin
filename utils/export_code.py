import os
from docx import Document

VERSION = "1.7"
SERVICE_NAME = "coin"
OUTPUT_NAME = f"{SERVICE_NAME}_export_{VERSION}.docx"

def export_code():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    print("[DEBUG] Đường dẫn file export_code.py:", current_dir)

    # 📁 Đường dẫn đến thư mục service (cùng cấp)
    service_path = os.path.abspath(os.path.join(current_dir, ".."))
    print("[DEBUG] Đường dẫn đến service:", service_path)

    if not os.path.exists(service_path):
        print("[❌] Không tìm thấy thư mục:", service_path)
        return

    # 📝 Tạo file docx
    doc = Document()
    doc.add_heading(f"📦 Mã nguồn: {SERVICE_NAME}", level=1)

    file_count = 0

    for root, dirs, files in os.walk(service_path):
        print("[DEBUG] Đang đọc thư mục:", root)

        for file in files:
            if not file.endswith(".py"):
                continue

            file_path = os.path.join(root, file)
            rel_path = os.path.relpath(file_path, service_path)

            print(f"  📄 Đọc file: {rel_path}")

            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                doc.add_heading(rel_path, level=2)
                doc.add_paragraph(content, style='Normal')
                file_count += 1

            except Exception as e:
                print(f"[⚠️] Không đọc được file: {file_path} → {type(e).__name__}: {str(e)}")

    # 📤 Lưu file
    output_path = os.path.abspath(os.path.join(current_dir, "../doc", OUTPUT_NAME))
    doc.save(output_path)
    print(f"✅ Hoàn tất. Đã ghi {file_count} file vào: {output_path}")

if __name__ == '__main__':
    export_code()
